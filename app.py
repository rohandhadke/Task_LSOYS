from flask import Flask, render_template, request, send_file
import openai
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

openai.api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)

@app.route("/", methods=['GET','POST'])
def index():
    return render_template("index.html")

@app.route("/generatereport", methods=['POST'])
def generatereport():
    try:
        project_name = request.form['project_title']
        project_tech = request.form['project_tech']
        project_abstract = request.form['project_abstract']
        project_introduction = request.form['project_introduction']
        project_problem_statement = request.form['project_problem_statement']
        project_new_things_implemented = request.form['project_new_implementation']
        
        sections = {
            "Introduction": project_introduction,
            "Problem Statement": project_problem_statement,
            "Scope of Project": generate_section("Scope of Project", project_abstract),
            "Literature Survey": generate_section("Literature Survey", project_abstract),
            "Methodology": generate_section("Methodology", project_abstract),
            "Project Life Cycle": generate_section("Project Life Cycle", project_abstract),
            "Timeline Chart": generate_section("Timeline Chart", project_abstract),
            "Time required for various stages": generate_section("Time required for various stages", project_abstract),
            "Design and working": generate_section("Design and working", project_abstract),
            "UML Diagrams": generate_section("UML Diagrams", project_abstract),
            "Testing of Project": generate_section("Testing of Project", project_abstract),
            "Cost Estimation": generate_section("Cost Estimation", project_abstract),
            "Result and Applications": generate_section("Result and Applications", project_abstract),
            "Coding": generate_section("Coding", project_abstract),
            "Conclusion": generate_section("Conclusion", project_abstract),
            "Evolution of project": generate_section("Evolution of project", project_abstract),
            "Memory Analysis": generate_section("Memory Analysis", project_abstract),
            "Platform": generate_section("Platform", project_abstract),
            "Performance": generate_section("Performance", project_abstract),
            "Advantage": generate_section("Advantage", project_abstract),
            "Disadvantage": generate_section("Disadvantage", project_abstract),
            "Future Scope": generate_section("Future Scope", project_abstract),
            "Bibliography": generate_section("Bibliography", project_abstract),
            "Book Reference": generate_section("Book Reference", project_abstract),
            "Web References": generate_section("Web References", project_abstract),
            "Appendix-B": generate_section("Appendix-B", project_abstract)
        }
        
        doc = create_docx(project_name, sections)
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(file_stream, as_attachment=True, download_name=f"{project_name}_report.docx")
    
    except Exception as e:
        app.logger.error(f"Error generating report: {e}")
        return {"error": "Failed to generate report"}, 500

def generate_section(title, context):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an assistant that helps generate project report sections."},
                {"role": "user", "content": f"Write a detailed section on '{title}' for a project report. Context: {context}"}
            ],
            max_tokens=500
        )
        return response.choices[0].message['content'].strip()
    
    except Exception as e:
        app.logger.error(f"Error generating section '{title}': {e}")
        return f"Error generating '{title}' section"

def create_docx(project_name, sections):
    doc = Document()

    title = doc.add_heading(project_name, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font_style(title, bold=True, size=16)

    for title, content in sections.items():
        heading = doc.add_heading(title, level=1)
        set_font_style(heading, bold=True, size=14)

        paragraph = doc.add_paragraph(content)
        set_font_style(paragraph, size=12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    return doc

def set_font_style(element, bold=False, size=12):
    for run in element.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0')
